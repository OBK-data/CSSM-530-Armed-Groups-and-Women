import pandas as pd
import nltk
#from wordcloud import WordCloud
import matplotlib.pyplot as plt
from nltk.stem import WordNetLemmatizer
import re
import string
from nltk.corpus import stopwords
from docx import Document
from docx.shared import Inches


lemmatizer = WordNetLemmatizer()
doc = Document()
def lemmatize_word(word):
    return lemmatizer.lemmatize(word)
org_name=["Hamas","Hezbollah","LTTE","IRA"]
#data = pd.read_csv("processed_news\IRA_parsed2.csv")
stops =  list(stopwords.words('english'))
removal_words=["would","one","two","three","four","five","six","seven","_","eight","nine","ten","'s","the","of","by","be","to","tiger",'"',",","eelam","-", "sinn fein","irish republican army","ltte","n't","tamil tigers","hamas","hezbollah","ira","irish republican army","NaN",",",".","'s",":",";","(",")","-","?","al","''","--","'"]
word_list=["she","her","herself","hers"]
stops.extend(removal_words)
#stops = [word for word in word_list if word not in stops]
rslt=pd.DataFrame()
data_small=pd.DataFrame()
def textfrequency(data):
    interest = ["woman","female"]
    data['lemma'] =data['lemma'].astype(str)
    data['head'] =data['head'].astype(str)
    data['lemma'] =data['lemma'].str.lower()
    data['head'] =data['head'].str.lower()
    data = data[~data['lemma'].isin(stops)]
    data = data[~data['head'].isin(stops)]
    data = data[~data['head'].str.isnumeric()]
    data = data[~data['lemma'].str.isnumeric()]
    data = data.dropna(subset=['head','lemma'])
    data["head"] = data["head"].apply(lemmatize_word)
    top_N = 50
    for place in range(len(interest)):
    # Apply lemmatization to the entire DataFrame
        data_small = data[data['lemma'].isin([interest[place]])]
        word_dist = nltk.FreqDist(data_small['head'])
        rslt=pd.DataFrame(word_dist.most_common(top_N),columns=['Word','Frequency'])
        doc.add_heading('Table - {} {} as head frequency'.format(org_name[num], interest[place]), level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Word'
        hdr_cells[1].text = 'Frequency'
        # Populate the table with DataFrame data
        for _, row in rslt.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Word'])
            row_cells[1].text = str(row['Frequency'])
        print('All frequencies')
        print('='*60)
        with pd.option_context('display.max_rows', None, 'display.max_columns', None):  # more options can be specified als
            print(rslt)
        print ('='*60)
        if interest[place]== "woman":
            data_small = data[data['head'].isin([interest[place]])]
            word_dist = nltk.FreqDist(data_small['lemma'])
            rslt=pd.DataFrame(word_dist.most_common(top_N),columns=['Word','Frequency'])
            doc.add_heading('Table - {} {} as dependent frequency'.format(org_name[num], interest[place]), level=1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Word'
            hdr_cells[1].text = 'Frequency'
            # Populate the table with DataFrame data
            for _, row in rslt.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Word'])
                row_cells[1].text = str(row['Frequency'])
            print('All frequencies')
            print('='*60)
            with pd.option_context('display.max_rows', None, 'display.max_columns', None):  # more options can be specified als
                print(rslt)
            print ('='*60)
        else:
            pass
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        # Add a table with header row
for num in range(len(org_name)):
    newslet= pd.read_csv("processed_news/{}_parsed2.csv".format(org_name[num]))
    textfrequency(newslet)
doc.save('table_of_frequency.docx')
